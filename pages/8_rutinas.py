import streamlit as st
from streamlit_gsheets import GSheetsConnection
import pandas as pd
from datetime import datetime, date, timedelta
import time
import random

# --- NUEVAS IMPORTACIONES PARA GENERAR WORD ---
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# ==========================================
# 1. CONFIGURACIÓN
# ==========================================
st.set_page_config(page_title="Sesiones de Entrenamiento", layout="centered")

# ==========================================
# 2. SEGURIDAD Y SESIÓN
# ==========================================
if "role" not in st.session_state or not st.session_state.role:
    st.switch_page("index.py")

rol = st.session_state.role
mi_id = st.session_state.user_id
mi_nombre = st.session_state.user_name

# ==========================================
# 3. CONEXIÓN A GOOGLE SHEETS
# ==========================================
conn = st.connection("gsheets", type=GSheetsConnection)

# ==========================================
# 4. FUNCIONES AUXILIARES Y GLOSARIO
# ==========================================

# --- NUEVA FUNCIÓN: GENERAR WORD 1/6 DE HOJA ---
def generar_word_sexto_hoja(entrenamiento_texto, fecha, titulo):
    doc = Document()
    
    # Configuración de márgenes para 1/6 de hoja (A4 aprox)
    # 4.5 pulgadas de margen izquierdo desplaza el texto a la derecha
    # 7.5 pulgadas de margen inferior asegura que no baje de la parte superior
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(7.5) 
    section.left_margin = Inches(4.5)   
    section.right_margin = Inches(0.5)
    
    # Título de la sesión
    p_titulo = doc.add_paragraph()
    run_t = p_titulo.add_run(f"{titulo} - {fecha}")
    run_t.bold = True
    run_t.font.size = Pt(10)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contenido del entrenamiento
    p_cont = doc.add_paragraph()
    # Limpiamos posibles saltos de línea excesivos
    texto_limpio = str(entrenamiento_texto).replace('\n\n', '\n')
    run_c = p_cont.add_run(texto_limpio)
    run_c.font.size = Pt(8.5) # Fuente pequeña para que entre en el recorte
    p_cont.alignment = WD_ALIGN_PARAGRAPH.LEFT

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def obtener_nombre_mes(n):
    meses = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", 
             "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    try:
        return meses[int(n)]
    except:
        return "Desconocido"

def get_periodo_mas_cercano(fecha_actual, registros_disponibles):
    """
    Calcula el período (año, mes) más cercano a la fecha actual basándose en los registros disponibles.
    Prioridad: 1. Actual, 2. Más cercano pasado, 3. Más cercano futuro.
    registros_disponibles: lista de tuplas [(anio, mes), ...]
    """
    if not registros_disponibles:
        return fecha_actual.year, fecha_actual.month

    # Convertir a objetos datetime para comparación (día 1)
    target = datetime(fecha_actual.year, fecha_actual.month, 1)
    fechas_disp = []
    
    # Filtrar y validar fechas
    for y, m in registros_disponibles:
        try:
            fechas_disp.append(datetime(int(y), int(m), 1))
        except:
            continue
            
    if not fechas_disp:
        return fecha_actual.year, fecha_actual.month

    # 1. Coincidencia exacta
    if target in fechas_disp:
        return target.year, target.month

    # 2. Buscar hacia atrás (Max del pasado)
    pasado = [d for d in fechas_disp if d < target]
    if pasado:
        mejor_pasado = max(pasado)
        return mejor_pasado.year, mejor_pasado.month

    # 3. Buscar hacia adelante (Min del futuro)
    futuro = [d for d in fechas_disp if d > target]
    if futuro:
        mejor_futuro = min(futuro)
        return mejor_futuro.year, mejor_futuro.month

    return fecha_actual.year, fecha_actual.month

def mostrar_referencias():
    with st.expander("📖 Glosario de Referencias (Clic para ver)"):
        st.markdown("""
        **INTENSIDADES**
        * **T (Tolerancia):** Intensidad alta 100 – 110%
        * **VC (Velocidad Corta):** Máxima velocidad
        * **VS (Velocidad Sostenida):** Mantener velocidad alta
        * **F1 (Vo2):** Intensidad 100%
        * **F2 (Super Aeróbico):** Intensidad 80-90%
        * **F3 (Sub Aeróbico):** Intensidad 70%
        * **Prog. (Progresivo):** De menor a mayor
        * **Reg (Regresivo):** De mayor a menor

        **ACTIVACIÓN Y ENTRADA EN CALOR**
        * **Ec (Entrada en Calor):** Nado suave inicial
        * **EcT (Ec Tensor):** Bíceps, Tríceps, Dorsales, Hombros, Pecho, Antebrazos
        * **EcM (Ec Movilidad):** Fuera del agua (Brazos, Cintura, Piernas, Tobillos, Cuello)
        * **Act (Activación):** Fuera del agua (Piernas, Brazos, Core)

        **TÉCNICA Y ESTILOS**
        * **B (Brazada):** C (Crol), E (Espalda), P (Pecho), M (Mariposa)
        * **Pat Ph:** Patada en Posición Hidrodinámica
        * **Pat Tabla:** Patada con tabla
        * **PB (Pull Brazada):** Uso de pullboy (c/e/p/m)
        * **CT:** Corrección Técnica
        
        **OTROS**
        * **m:** Metros
        * **p:** Pausa estática
        * **p act:** Pausa Activa
        * **D/:** Dentro del tiempo
        * **C/:** Con tiempo de pausa
        """)

# --- MANEJO DE RETRYS PARA API ---
def actualizar_con_retry(worksheet, data, max_retries=5):
    for i in range(max_retries):
        try:
            conn.update(worksheet=worksheet, data=data)
            return True, None 
        except Exception as e:
            error_msg = str(e).lower()
            if "429" in error_msg or "quota" in error_msg or "rate limit" in error_msg:
                wait_time = (2 ** i) + random.uniform(0, 1)
                time.sleep(wait_time)
                continue 
            else:
                return False, e
    return False, "Tiempo de espera agotado (API ocupada)."

# --- LECTURA DE DATOS (CON CACHÉ) ---
@st.cache_data(ttl="5s")
def cargar_datos_rutinas_view():
    try:
        try:
            df_rut = conn.read(worksheet="Rutinas").copy()
        except:
            df_rut = pd.DataFrame(columns=["id_rutina", "anio_rutina", "mes_rutina", "nro_sesion", "texto_rutina"])
        
        try:
            df_seg = conn.read(worksheet="Rutinas_Seguimiento").copy()
        except:
            df_seg = pd.DataFrame(columns=["id_rutina", "codnadador", "fecha_realizada"])

        try:
            df_nad = conn.read(worksheet="Nadadores").copy()
        except:
            df_nad = pd.DataFrame(columns=["codnadador", "nombre", "apellido"])
            
        # Normalización de tipos
        if not df_rut.empty:
            df_rut['anio_rutina'] = pd.to_numeric(df_rut['anio_rutina'], errors='coerce').fillna(0).astype(int)
            df_rut['mes_rutina'] = pd.to_numeric(df_rut['mes_rutina'], errors='coerce').fillna(0).astype(int)
            df_rut['nro_sesion'] = pd.to_numeric(df_rut['nro_sesion'], errors='coerce').fillna(0).astype(int)
            
        if not df_seg.empty:
            df_seg['codnadador'] = pd.to_numeric(df_seg['codnadador'], errors='coerce').fillna(0).astype(int)

        if not df_nad.empty:
             df_nad['codnadador'] = pd.to_numeric(df_nad['codnadador'], errors='coerce').fillna(0).astype(int)
            
        return df_rut, df_seg, df_nad
    except Exception as e:
        st.error(f"Error visual al cargar datos: {e}")
        return None, None, None

def leer_dataset_fresco(worksheet):
    try:
        df = conn.read(worksheet=worksheet, ttl=0).copy()
        return df
    except Exception as e:
        st.error(f"Error de conexión fresca con {worksheet}: {e}")
        return None

def calcular_proxima_sesion(df, anio, mes):
    if df is None or df.empty: return 1
    filtro = df[(df['anio_rutina'] == anio) & (df['mes_rutina'] == mes)]
    if filtro.empty: return 1
    return int(filtro['nro_sesion'].max()) + 1

# ==========================================
# 5. FUNCIONES DE ESCRITURA
# ==========================================

def guardar_seguimiento(id_rutina, id_nadador):
    df_seg = leer_dataset_fresco("Rutinas_Seguimiento")
    if df_seg is None: return False 
    
    if not df_seg.empty:
        df_seg['codnadador'] = pd.to_numeric(df_seg['codnadador'], errors='coerce').fillna(0).astype(int)

    existe = df_seg[(df_seg['id_rutina'] == id_rutina) & (df_seg['codnadador'] == id_nadador)]
    
    if existe.empty:
        hora_arg = datetime.now() - timedelta(hours=3)
        nuevo_registro = pd.DataFrame([{
            "id_rutina": id_rutina,
            "codnadador": id_nadador,
            "fecha_realizada": hora_arg.strftime("%Y-%m-%d %H:%M:%S")
        }])
        df_final = pd.concat([df_seg, nuevo_registro], ignore_index=True)
        
        exito, error = actualizar_con_retry("Rutinas_Seguimiento", df_final)
        if exito:
            st.cache_data.clear() 
            return True
        else:
            st.error(f"Error al guardar: {error}")
            return False
    return False

def borrar_seguimiento(id_rutina, id_nadador):
    df_seg = leer_dataset_fresco("Rutinas_Seguimiento")
    if df_seg is None: return False
    
    if not df_seg.empty:
        df_seg['codnadador'] = pd.to_numeric(df_seg['codnadador'], errors='coerce').fillna(0).astype(int)

    df_final = df_seg[~((df_seg['id_rutina'] == id_rutina) & (df_seg['codnadador'] == id_nadador))]
    
    exito, error = actualizar_con_retry("Rutinas_Seguimiento", df_final)
    if exito:
        st.cache_data.clear()
        return True
    else:
        st.error(f"Error al borrar: {error}")
        return False

def eliminar_sesion_admin(id_rutina):
    df_rut = leer_dataset_fresco("Rutinas")
    if df_rut is None: return "❌ Error de conexión."
    
    rutina_a_borrar = df_rut[df_rut['id_rutina'] == id_rutina]
    
    if rutina_a_borrar.empty:
        return "⚠️ La sesión no existe."
    
    r_anio = rutina_a_borrar.iloc[0]['anio_rutina']
    r_mes = rutina_a_borrar.iloc[0]['mes_rutina']
    r_sesion = rutina_a_borrar.iloc[0]['nro_sesion']
    
    rutinas_mes = df_rut[(df_rut['anio_rutina'] == r_anio) & (df_rut['mes_rutina'] == r_mes)]
    max_sesion = rutinas_mes['nro_sesion'].max()
    
    if r_sesion < max_sesion:
        return f"🚫 Solo se permite eliminar la última sesión del mes (Sesión {max_sesion})."

    df_rut_final = df_rut[df_rut['id_rutina'] != id_rutina]
    
    exito, error = actualizar_con_retry("Rutinas", df_rut_final)
    
    if exito:
        st.cache_data.clear()
        return "🗑️ Sesión eliminada correctamente."
    else:
        return f"❌ Error al eliminar: {error}"

def guardar_sesion_admin(anio, mes, sesion, texto):
    df_rut = leer_dataset_fresco("Rutinas")
    if df_rut is None: return "❌ Error CRÍTICO de conexión."

    if not df_rut.empty:
        df_rut['anio_rutina'] = pd.to_numeric(df_rut['anio_rutina'], errors='coerce').fillna(0).astype(int)
        df_rut['mes_rutina'] = pd.to_numeric(df_rut['mes_rutina'], errors='coerce').fillna(0).astype(int)
        df_rut['nro_sesion'] = pd.to_numeric(df_rut['nro_sesion'], errors='coerce').fillna(0).astype(int)

    nuevo_id = f"{anio}-{mes:02d}-S{sesion:02d}"
    mask = df_rut['id_rutina'] == nuevo_id
    
    nueva_fila = {
        "id_rutina": nuevo_id,
        "anio_rutina": int(anio),
        "mes_rutina": int(mes),
        "nro_sesion": int(sesion),
        "texto_rutina": texto
    }
    
    if df_rut[mask].empty:
        df_rut = pd.concat([df_rut, pd.DataFrame([nueva_fila])], ignore_index=True)
        msg = "✅ Sesión creada correctamente."
    else:
        df_rut.loc[mask, "texto_rutina"] = texto
        msg = "✅ Sesión actualizada correctamente."
        
    exito, error = actualizar_con_retry("Rutinas", df_rut)
    if exito:
        st.cache_data.clear()
    return msg if exito else f"❌ Error al escribir: {error}"

def activar_calculo_auto():
    st.session_state.trigger_calculo = True

# ==========================================
# 6. COMPONENTES VISUALES
# ==========================================

def render_tarjeta_individual(row, df_seg, key_suffix):
    r_id = row['id_rutina']
    r_sesion = row['nro_sesion']
    r_texto = row['texto_rutina']
    r_anio = row.get('anio_rutina', '')
    r_mes = row.get('mes_rutina', '')
    
    check = df_seg[(df_seg['id_rutina'] == r_id) & (df_seg['codnadador'] == mi_id)]
    esta_realizada = not check.empty
    fecha_str = ""
    if esta_realizada:
        fecha_obj = pd.to_datetime(check.iloc[0]['fecha_realizada'])
        fecha_str = fecha_obj.strftime("%d/%m")

    if esta_realizada:
        borde = "#2E7D32" # Verde
        bg = "#1B2E1B"
    else:
        borde = "#444" # Gris
        bg = "#262730"
    
    with st.container():
        st.markdown(f"""<div style="border: 2px solid {borde}; border-radius: 10px; background-color: {bg}; padding: 15px; margin-bottom: 15px;">""", unsafe_allow_html=True)
        
        if esta_realizada:
            c_head, c_act = st.columns([8, 1])
            with c_head:
                st.markdown(f"#### ✅ Sesión {r_sesion} <span style='font-size:14px; color:#888'>({fecha_str})</span>", unsafe_allow_html=True)
                st.markdown(f"<div style='text-decoration: line-through; color: #aaa;'>", unsafe_allow_html=True)
                st.markdown(r_texto)
                st.markdown("</div>", unsafe_allow_html=True)
                
                # Botón de descarga para Word 1/6 hoja
                doc_buffer = generar_word_sexto_hoja(r_texto, f"{r_mes:02d}/{r_anio}", f"SESIÓN {r_sesion}")
                st.download_button(
                    label="📄 Descargar para pileta (1/6 hoja)",
                    data=doc_buffer,
                    file_name=f"Rutina_S{r_sesion}_{r_mes}_{r_anio}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{r_id}_{key_suffix}_done"
                )
            with c_act:
                if st.button("❌", key=f"un_{r_id}_{key_suffix}", help="Desmarcar"):
                    with st.spinner("Desmarcando..."):
                        borrar_seguimiento(r_id, mi_id)
                    st.rerun()
        else:
            c_head, c_act = st.columns([5, 2])
            with c_head:
                st.markdown(f"#### ⭕ Sesión {r_sesion}")
                st.markdown(r_texto)
                
                # Botón de descarga para Word 1/6 hoja
                doc_buffer = generar_word_sexto_hoja(r_texto, f"{r_mes:02d}/{r_anio}", f"SESIÓN {r_sesion}")
                st.download_button(
                    label="📄 Descargar para pileta (1/6 hoja)",
                    data=doc_buffer,
                    file_name=f"Rutina_S{r_sesion}_{r_mes}_{r_anio}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{r_id}_{key_suffix}_pend"
                )
            with c_act:
                st.write("") 
                if st.button("🏊 DÍA GANADO", key=f"do_{r_id}_{key_suffix}", type="primary"):
                    with st.spinner("Guardando..."):
                        guardar_seguimiento(r_id, mi_id)
                    st.rerun()
        
        st.markdown("</div>", unsafe_allow_html=True)

def render_feed_activo(df_rut, df_seg, anio_ver, mes_ver, key_suffix=""):
    rutinas_filtradas = df_rut[
        (df_rut['anio_rutina'] == anio_ver) & 
        (df_rut['mes_rutina'] == mes_ver)
    ].copy()

    if rutinas_filtradas.empty:
        st.info(f"No hay sesiones cargadas para {obtener_nombre_mes(mes_ver)} {anio_ver}.")
        return

    rutinas_filtradas.sort_values(by='nro_sesion', ascending=True, inplace=True)
    mostrar_referencias()

    l_pendientes = []
    l_completadas = []

    for index, row in rutinas_filtradas.iterrows():
        check = df_seg[(df_seg['id_rutina'] == row['id_rutina']) & (df_seg['codnadador'] == mi_id)]
        if not check.empty:
            l_completadas.append(row)
        else:
            l_pendientes.append(row)

    if l_completadas:
        with st.expander(f"✅ Historial Completado ({len(l_completadas)})", expanded=False):
            for row in reversed(l_completadas):
                render_tarjeta_individual(row, df_seg, key_suffix)
        st.write("---")

    if l_pendientes:
        st.markdown("#### 🚀 Próximas Sesiones")
        for row in l_pendientes:
            render_tarjeta_individual(row, df_seg, key_suffix)
    else:
        if l_completadas:
            st.success("¡Excelente! Has completado todas las sesiones del mes. 🏆")

# --- HISTORIAL VISUAL MEJORADO ---
def render_historial_compacto(df_rut, df_seg, anio, mes, id_usuario_objetivo):
    rutinas_mes = df_rut[
        (df_rut['anio_rutina'] == anio) & 
        (df_rut['mes_rutina'] == mes)
    ].sort_values('nro_sesion')

    if rutinas_mes.empty:
        st.info("No hay planificación cargada para este mes.")
        return

    total_rutinas = len(rutinas_mes)
    completadas = 0
    lista_render = []

    for _, r in rutinas_mes.iterrows():
        r_id = r['id_rutina']
        check = df_seg[(df_seg['id_rutina'] == r_id) & (df_seg['codnadador'] == id_usuario_objetivo)]
        hecho = not check.empty
        fecha_txt = "-"
        
        if hecho:
            completadas += 1
            fecha_obj = pd.to_datetime(check.iloc[0]['fecha_realizada'])
            fecha_txt = fecha_obj.strftime("%d/%m")
        
        lista_render.append({
            "nro": r['nro_sesion'],
            "hecho": hecho,
            "fecha": fecha_txt
        })

    porcentaje = int((completadas / total_rutinas) * 100) if total_rutinas > 0 else 0
    nombre_mes = obtener_nombre_mes(mes).upper()
    
    # -------------------------------------------------------------
    # EFECTO DOPAMINA: Tarjeta de Logro al 100% (Solo Nadador)
    # -------------------------------------------------------------
    if st.session_state.role == "N" and porcentaje == 100:
        celeb_key = f"celeb_{anio}_{mes}"
        if celeb_key not in st.session_state:
            st.balloons()
            st.session_state[celeb_key] = True
            
        st.markdown(f"""
        <div style="
            background: linear-gradient(135deg, #E30613 0%, #000000 100%);
            border-radius: 12px;
            padding: 20px;
            margin-bottom: 25px;
            text-align: center;
            border: 2px solid #fff;
            box-shadow: 0 4px 15px rgba(227, 6, 19, 0.4);
        ">
            <div style="font-size: 40px; margin-bottom: 10px;">🏆</div>
            <h2 style="color: white; margin: 0; font-weight: 800; text-transform: uppercase; font-style: italic;">¡RUTINA COMPLETADA!</h2>
            <p style="color: #f0f0f0; font-size: 16px; margin-top: 5px; font-weight: 500;">
                Lograste el 100% de asistencia en {nombre_mes}.<br>
                <span style="font-size: 14px; opacity: 0.9; font-style: italic;">"Sumás constancia, sumás mejora"</span>
            </p>
        </div>
        """, unsafe_allow_html=True)
    # -------------------------------------------------------------

    # SCORECARD DIVIDIDO
    st.markdown(f"""
    <div style="background-color: #262730; border-radius: 8px; padding: 20px; border: 1px solid #444; margin-bottom: 20px;">
        <div style="color: #aaa; font-size: 13px; font-weight: bold; margin-bottom: 12px; letter-spacing: 1px;">
            ASISTENCIA {nombre_mes}
        </div>
        <div style="display: flex; justify-content: space-between; align-items: center;">
            <div>
                <span style="font-size: 42px; font-weight: bold; color: white;">{porcentaje}%</span>
            </div>
            <div style="text-align: right;">
                <div style="font-size: 24px; font-weight: bold; color: #ddd; line-height: 1;">
                    {completadas} <span style="font-size: 16px; color: #888;">de</span> {total_rutinas}
                </div>
                <div style="font-size: 11px; color: #888; margin-top: 4px; letter-spacing: 0.5px;">SESIONES COMPLETADAS</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.progress(porcentaje / 100)
    st.write("") 

    # LISTADO DE SESIONES
    st.caption(f"DETALLE DE ACTIVIDAD")

    for item in lista_render:
        if item['hecho']:
            color = "#2E7D32" # Verde
            icono = "✅"
            texto_estado = "NADADO"
            bg_badge = "#1B5E20"
        else:
            color = "#555" # Gris
            icono = "⭕"
            texto_estado = "PENDIENTE"
            bg_badge = "#444"

        st.markdown(f"""
        <div style="
            background-color: #262730;
            border-radius: 6px;
            padding: 10px 15px;
            margin-bottom: 8px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            border-left: 4px solid {color};
        ">
            <div style="display: flex; align-items: center; gap: 10px;">
                <span style="font-size: 16px;">{icono}</span>
                <span style="font-weight: 500; color: white;">Sesión {item['nro']}</span>
            </div>
            <div style="text-align: right;">
                <span style="
                    background-color: {bg_badge};
                    color: white;
                    padding: 3px 8px;
                    border-radius: 4px;
                    font-size: 10px;
                    font-weight: bold;
                    letter-spacing: 0.5px;
                ">{texto_estado}</span>
                <div style="font-size: 11px; color: #aaa; margin-top: 3px;">{item['fecha']}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

# ==========================================
# 7. LÓGICA PRINCIPAL (MAIN)
# ==========================================

df_rutinas, df_seguimiento, df_nadadores = cargar_datos_rutinas_view()

# --- INTERFAZ DE USUARIO ---
st.title("📝 Sesiones de Entrenamiento")
st.subheader(f"{mi_nombre}")

if df_rutinas is None:
    st.error("No se pudieron cargar los datos. Verifica tu conexión.")
    st.stop()

st.write("---")

# ==========================
# VISTA: PROFESOR (M) / ADMIN
# ==========================
if rol in ["M", "P"]:
    
    if "g_anio" not in st.session_state: st.session_state.g_anio = datetime.now().year
    if "g_mes" not in st.session_state: st.session_state.g_mes = datetime.now().month

    # Trigger de cálculo inicial
    if st.session_state.get("trigger_calculo", False) or "admin_sesion" not in st.session_state:
        if df_rutinas is not None:
            prox = calcular_proxima_sesion(df_rutinas, st.session_state.g_anio, st.session_state.g_mes)
            st.session_state.admin_sesion = min(prox, 31)
        else:
            st.session_state.admin_sesion = 1
        st.session_state.trigger_calculo = False

    # Lógica de avance manual seguro
    if st.session_state.get("forzar_avance_sesion", False):
        if "admin_sesion" in st.session_state and st.session_state.admin_sesion < 31:
             st.session_state.admin_sesion += 1
        st.session_state.forzar_avance_sesion = False

    with st.expander("⚙️ Gestión de Sesiones (Crear/Editar)", expanded=False):
        st.markdown("##### Editor de Sesiones")
        
        c1, c2, c3 = st.columns([1, 1, 1])
        
        anio_actual = datetime.now().year
        anios_gest = sorted(list(set(df_rutinas['anio_rutina'].unique().tolist() + [anio_actual])), reverse=True)
        meses_indices = list(range(1, 13))
        mapa_meses = {i: obtener_nombre_mes(i) for i in meses_indices}

        with c1: 
            st.number_input("Año", min_value=2020, max_value=2030, key="g_anio", on_change=activar_calculo_auto)
        with c2: 
            st.selectbox("Mes", meses_indices, format_func=lambda x: mapa_meses[x], key="g_mes", on_change=activar_calculo_auto)
        with c3: 
            st.number_input("Nro Sesión", min_value=1, max_value=31, key="admin_sesion")
            
        id_busqueda = f"{st.session_state.g_anio}-{st.session_state.g_mes:02d}-S{st.session_state.admin_sesion:02d}"
        row_existente = df_rutinas[df_rutinas['id_rutina'] == id_busqueda]
        
        es_edicion = not row_existente.empty
        texto_previo = row_existente.iloc[0]['texto_rutina'] if es_edicion else ""
        
        estado_txt = "✏️ Editando Existente" if es_edicion else "✨ Nueva Sesión"
        st.caption(f"ID: {id_busqueda} | {estado_txt}")
        
        with st.form("form_rutina"):
            f_texto = st.text_area("Contenido", value=texto_previo, height=200, key=f"editor_texto_{id_busqueda}")
            
            if es_edicion:
                c_del, c_save = st.columns([1, 2])
                with c_del:
                    delete_btn = st.form_submit_button("🗑️ Eliminar Sesión", type="secondary")
                with c_save:
                    submitted = st.form_submit_button("💾 Actualizar Sesión", type="primary")
            else:
                submitted = st.form_submit_button("💾 Crear Sesión", type="primary")
                delete_btn = False

            if submitted:
                if f_texto.strip() == "":
                    st.error("Texto vacío.")
                else:
                    msg = guardar_sesion_admin(st.session_state.g_anio, st.session_state.g_mes, st.session_state.admin_sesion, f_texto)
                    if "Error" in msg:
                        st.error(msg)
                    else:
                        st.success(msg)
                        st.session_state.forzar_avance_sesion = True
                        time.sleep(0.5)
                        st.rerun()
            
            if delete_btn:
                msg = eliminar_sesion_admin(id_busqueda)
                if "Error" in msg:
                    st.error(msg)
                else:
                    st.warning(msg)
                    st.session_state.trigger_calculo = True
                    time.sleep(1)
                    st.rerun()

    st.divider()

    tab_explorar, tab_seguimiento = st.tabs(["📖 Explorar Sesiones", "📊 Seguimiento Alumnos"])
    
    # --- LOGICA DE DEFAULT (PERFIL M - EXPLORAR) ---
    unique_periods_global = df_rutinas[['anio_rutina', 'mes_rutina']].drop_duplicates().values.tolist()
    def_y_ex, def_m_ex = get_periodo_mas_cercano(datetime.now(), unique_periods_global)

    with tab_explorar:
        col_v1, col_v2 = st.columns(2)
        v_anios = sorted(list(set(df_rutinas['anio_rutina'].unique().tolist() + [datetime.now().year])), reverse=True)
        
        idx_y_ex = v_anios.index(def_y_ex) if def_y_ex in v_anios else 0
        with col_v1: sel_a = st.selectbox("Año", v_anios, index=idx_y_ex, key="adm_v_a")
        
        meses_disp_explorar = sorted(df_rutinas[df_rutinas['anio_rutina'] == sel_a]['mes_rutina'].unique().tolist())
        
        if not meses_disp_explorar:
             with col_v2: st.warning("Sin datos")
             sel_m = None
        else:
             mapa_meses_ex = {i: obtener_nombre_mes(i) for i in meses_disp_explorar}
             
             # Intentar mantener el mes por defecto si existe en la lista filtrada
             idx_m_ex = 0
             if def_m_ex in meses_disp_explorar:
                 idx_m_ex = meses_disp_explorar.index(def_m_ex)
             elif meses_disp_explorar:
                 # Si el mes por defecto no está (ej: cambiamos de año), buscamos el más cercano en la lista
                 target_dt = datetime(sel_a, def_m_ex, 1)
                 avail_dt = [datetime(sel_a, m, 1) for m in meses_disp_explorar]
                 # Reutilizamos logica simple: mas cercano
                 closest = min(avail_dt, key=lambda x: abs(x - target_dt))
                 idx_m_ex = meses_disp_explorar.index(closest.month)

             with col_v2: sel_m = st.selectbox("Mes", meses_disp_explorar, format_func=lambda x: mapa_meses_ex[x], index=idx_m_ex, key="adm_v_m")
        
        if sel_m:
            render_feed_activo(df_rutinas, df_seguimiento, sel_a, sel_m, key_suffix="admin_view")

    # --- LOGICA DE DEFAULT (PERFIL M - SEGUIMIENTO) ---
    # Usamos la misma logica global para los defaults de tiempo
    with tab_seguimiento:
        st.info("Seleccione un alumno para ver su cumplimiento histórico.")
        
        ids_activos = df_seguimiento['codnadador'].unique()
        df_nad_activos = df_nadadores[df_nadadores['codnadador'].isin(ids_activos)].copy()
        
        if df_nad_activos.empty:
            st.warning("⚠️ Aún no hay alumnos con sesiones completadas.")
        else:
            df_nad_activos['NombreCompleto'] = df_nad_activos['apellido'] + ", " + df_nad_activos['nombre']
            lista_nads = df_nad_activos[['codnadador', 'NombreCompleto']].sort_values('NombreCompleto')
            
            col_s1, col_s2, col_s3 = st.columns([2, 1, 1])
            with col_s1:
                sel_nad_id = st.selectbox(
                    "Alumno", 
                    lista_nads['codnadador'], 
                    format_func=lambda x: lista_nads[lista_nads['codnadador'] == x]['NombreCompleto'].values[0]
                )
            
            # Aplicamos defaults calculados previamente (Global Best Period)
            # Para Año
            idx_y_seg = v_anios.index(def_y_ex) if def_y_ex in v_anios else 0
            with col_s2:
                sel_a_seg = st.selectbox("Año", v_anios, index=idx_y_seg, key="seg_a")
            
            meses_disp_seg = sorted(df_rutinas[df_rutinas['anio_rutina'] == sel_a_seg]['mes_rutina'].unique().tolist())
            
            if not meses_disp_seg:
                 with col_s3: st.warning("Sin datos")
                 sel_m_seg = None
            else:
                 mapa_meses_seg = {i: obtener_nombre_mes(i) for i in meses_disp_seg}
                 
                 # Lógica de índice para mes
                 idx_m_seg = 0
                 if def_m_ex in meses_disp_seg:
                     idx_m_seg = meses_disp_seg.index(def_m_ex)
                 elif meses_disp_seg:
                     # Fallback al más cercano
                     target_dt = datetime(sel_a_seg, def_m_ex, 1)
                     avail_dt = [datetime(sel_a_seg, m, 1) for m in meses_disp_seg]
                     closest = min(avail_dt, key=lambda x: abs(x - target_dt))
                     idx_m_seg = meses_disp_seg.index(closest.month)

                 with col_s3: sel_m_seg = st.selectbox("Mes", meses_disp_seg, format_func=lambda x: mapa_meses_seg[x], index=idx_m_seg, key="seg_m")
                
            st.markdown(f"**Reporte para:** {lista_nads[lista_nads['codnadador']==sel_nad_id]['NombreCompleto'].values[0]}")
            if sel_m_seg:
                render_historial_compacto(df_rutinas, df_seguimiento, sel_a_seg, sel_m_seg, sel_nad_id)

# ==========================
# VISTA: NADADOR (N)
# ==========================
else:
    tab_curso, tab_hist = st.tabs(["📅 Mes en Curso", "📜 Historial"])
    
    with tab_curso:
        hoy = datetime.now()
        
        st.markdown(f"### Sesiones de {obtener_nombre_mes(hoy.month)} {hoy.year}")
        render_feed_activo(df_rutinas, df_seguimiento, hoy.year, hoy.month, key_suffix="nad_curso")
        
    with tab_hist:
        st.markdown("#### Registro de Cumplimiento")
        
        # --- LOGICA DE DEFAULT (PERFIL N - HISTORIAL) ---
        # 1. Obtener períodos disponibles GLOBALMENTE en rutinas
        unique_periods_n = df_rutinas[['anio_rutina', 'mes_rutina']].drop_duplicates().values.tolist()
        
        # 2. Calcular mejor fecha
        def_y_n, def_m_n = get_periodo_mas_cercano(datetime.now(), unique_periods_n)

        c_h1, c_h2 = st.columns(2)
        anios_disp = sorted(list(set(df_rutinas['anio_rutina'].unique())), reverse=True)
        if not anios_disp: anios_disp = [datetime.now().year]
        
        # Índice año default
        idx_y_n = anios_disp.index(def_y_n) if def_y_n in anios_disp else 0
        
        with c_h1: h_anio = st.selectbox("Año", anios_disp, index=idx_y_n, key="h_a")
        
        meses_en_anio = sorted(df_rutinas[df_rutinas['anio_rutina'] == h_anio]['mes_rutina'].unique().tolist())
        
        if not meses_en_anio:
            with c_h2: 
                st.warning("Sin datos.")
                h_mes = None
        else:
            mapa_meses = {i: obtener_nombre_mes(i) for i in meses_en_anio}
            
            # Índice mes default
            idx_m_n = 0
            if def_m_n in meses_en_anio:
                idx_m_n = meses_en_anio.index(def_m_n)
            elif meses_en_anio:
                 # Fallback inteligente
                 target_dt = datetime(h_anio, def_m_n, 1)
                 avail_dt = [datetime(h_anio, m, 1) for m in meses_en_anio]
                 closest = min(avail_dt, key=lambda x: abs(x - target_dt))
                 idx_m_n = meses_en_anio.index(closest.month)

            with c_h2: 
                h_mes = st.selectbox("Mes", meses_en_anio, format_func=lambda x: mapa_meses[x], index=idx_m_n, key="h_m")
            
        if h_mes:
            render_historial_compacto(df_rutinas, df_seguimiento, h_anio, h_mes, mi_id)
